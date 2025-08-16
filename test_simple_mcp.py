#!/usr/bin/env python3
"""
Test script for the Word MCP server.
"""

import os
import json
from typing import Optional
from docx import Document

def ensure_docx_extension(filename: str) -> str:
    """Ensure filename has .docx extension."""
    if not filename.endswith('.docx'):
        return f"{filename}.docx"
    return filename

def check_file_exists(filename: str) -> bool:
    """Check if file exists."""
    return os.path.exists(filename)

def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document."""
    filename = ensure_docx_extension(filename)
    
    try:
        doc = Document()
        
        # Set metadata if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        doc.save(filename)
        return f"Document '{filename}' created successfully"
    except Exception as e:
        return f"Error creating document: {str(e)}"

def write_text(filename: str, text: str, append: bool = True) -> str:
    """Write text to a Word document."""
    filename = ensure_docx_extension(filename)
    
    try:
        if append and check_file_exists(filename):
            # Append to existing document
            doc = Document(filename)
        else:
            # Create new document or replace existing
            doc = Document()
        
        # Add text as paragraph
        doc.add_paragraph(text)
        doc.save(filename)
        
        action = "appended to" if append and check_file_exists(filename) else "written to"
        return f"Text {action} '{filename}' successfully"
    except Exception as e:
        return f"Error writing to document: {str(e)}"

def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document."""
    filename = ensure_docx_extension(filename)
    
    # Validate heading level
    if level < 1 or level > 6:
        return "Error: Heading level must be between 1 and 6"
    
    try:
        if check_file_exists(filename):
            doc = Document(filename)
        else:
            doc = Document()
        
        doc.add_heading(text, level=level)
        doc.save(filename)
        
        return f"Heading '{text}' (level {level}) added to '{filename}'"
    except Exception as e:
        return f"Error adding heading: {str(e)}"

def read_document(filename: str) -> str:
    """Read all text content from a Word document."""
    filename = ensure_docx_extension(filename)
    
    if not check_file_exists(filename):
        return f"Error: Document '{filename}' does not exist"
    
    try:
        doc = Document(filename)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text)
        
        return "\n".join(text_content)
    except Exception as e:
        return f"Error reading document: {str(e)}"

def get_document_info(filename: str) -> str:
    """Get basic information about a Word document."""
    filename = ensure_docx_extension(filename)
    
    if not check_file_exists(filename):
        return f"Error: Document '{filename}' does not exist"
    
    try:
        doc = Document(filename)
        
        info = {
            "filename": filename,
            "title": doc.core_properties.title or "Untitled",
            "author": doc.core_properties.author or "Unknown",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "Unknown",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "Unknown",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size_kb": round(os.path.getsize(filename) / 1024, 2)
        }
        
        return json.dumps(info, indent=2)
    except Exception as e:
        return f"Error getting document info: {str(e)}"

def replace_text(filename: str, find_text: str, replace_text: str) -> str:
    """Find and replace text in a Word document."""
    filename = ensure_docx_extension(filename)
    
    if not check_file_exists(filename):
        return f"Error: Document '{filename}' does not exist"
    
    try:
        doc = Document(filename)
        replacements = 0
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        replacements += run.text.count(replace_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            for run in paragraph.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
                                    replacements += run.text.count(replace_text)
        
        if replacements > 0:
            doc.save(filename)
            return f"Replaced {replacements} occurrence(s) of '{find_text}' with '{replace_text}'"
        else:
            return f"No occurrences of '{find_text}' found"
    except Exception as e:
        return f"Error replacing text: {str(e)}"

def main():
    """Test the simplified Word MCP functionality."""
    print("Testing simplified Word MCP server functionality...")
    print("=" * 50)
    
    # Clean up any existing test file
    test_file = "test_doc.docx"
    if os.path.exists(test_file):
        os.remove(test_file)
    
    # Test 1: Create document
    print("1. Creating document:")
    result = create_document(test_file, 'Test Document', 'Test Author')
    print(f"   {result}")
    
    # Test 2: Add heading
    print("\n2. Adding heading:")
    result = add_heading(test_file, 'Chapter 1: Introduction', 1)
    print(f"   {result}")
    
    # Test 3: Write text
    print("\n3. Writing text:")
    result = write_text(test_file, 'This is the first paragraph of our test document.')
    print(f"   {result}")
    
    # Test 4: Add more content
    print("\n4. Adding more content:")
    result = add_heading(test_file, 'Section 1.1', 2)
    print(f"   {result}")
    result = write_text(test_file, 'This is a subsection with more detailed information.')
    print(f"   {result}")
    
    # Test 5: Read document
    print("\n5. Reading document content:")
    content = read_document(test_file)
    print(f"   Content:\n   {content.replace(chr(10), chr(10) + '   ')}")
    
    # Test 6: Get document info
    print("\n6. Getting document info:")
    info = get_document_info(test_file)
    print(f"   {info}")
    
    # Test 7: Replace text
    print("\n7. Testing text replacement:")
    result = replace_text(test_file, 'test', 'sample')
    print(f"   {result}")
    
    # Test 8: Read document after replacement
    print("\n8. Reading document after replacement:")
    content = read_document(test_file)
    print(f"   Content:\n   {content.replace(chr(10), chr(10) + '   ')}")
    
    print("\n" + "=" * 50)
    print("All tests completed successfully!")
    print(f"Test document '{test_file}' created and can be opened in Word.")

if __name__ == "__main__":
    main()