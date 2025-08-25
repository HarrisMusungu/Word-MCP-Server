#!/usr/bin/env python3
"""
Simplified Word Document MCP Server using FastMCP.
Provides only essential Read and Write operations for Word documents.
"""

import os
import json
import subprocess
from typing import Optional, List
from docx import Document
from fastmcp import FastMCP

# Initialize FastMCP server
mcp = FastMCP("Simple Word Document Server")

def ensure_docx_extension(filename: str) -> str:
    """Ensure filename has .docx extension."""
    if not filename.endswith('.docx'):
        return f"{filename}.docx"
    return filename

def check_file_exists(filename: str) -> bool:
    """Check if file exists."""
    return os.path.exists(filename)

# READ OPERATIONS

@mcp.tool()
def read_document(filename: str) -> str:
    """Read all text content from a Word document.
    
    Args:
        filename: Path to the Word document (.docx extension will be added if missing)
    
    Returns:
        String containing all text from the document
    """
    filename = ensure_docx_extension(filename)
    
    if not check_file_exists(filename):
        return f"Error: Document '{filename}' does not exist"
    
    try:
        doc = Document(filename)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():  # Only add non-empty paragraphs
                text_content.append(paragraph.text)
        
        return "\n".join(text_content)
    except Exception as e:
        return f"Error reading document: {str(e)}"

@mcp.tool()
def get_document_info(filename: str) -> str:
    """Get basic information about a Word document.
    
    Args:
        filename: Path to the Word document
    
    Returns:
        JSON string with document metadata
    """
    filename = ensure_docx_extension(filename)
    
    if not check_file_exists(filename):
        return f"Error: Document '{filename}' does not exist"
    
    try:
        doc = Document(filename)
        
        info = {
            "filename": filename,
            "title": doc.core_properties.title or "Untitled",
            "author": doc.core_properties.author or "Unknown",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "Unknown",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "Unknown",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "file_size_kb": round(os.path.getsize(filename) / 1024, 2)
        }
        
        return json.dumps(info, indent=2)
    except Exception as e:
        return f"Error getting document info: {str(e)}"

@mcp.tool()
def list_documents(directory: str = ".") -> str:
    """List all Word documents in a directory.
    
    Args:
        directory: Directory path to search (defaults to current directory)
    
    Returns:
        List of .docx files found
    """
    try:
        if not os.path.exists(directory):
            return f"Error: Directory '{directory}' does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in '{directory}'"
        
        result = f"Found {len(docx_files)} Word documents in '{directory}':\n"
        for file in sorted(docx_files):
            file_path = os.path.join(directory, file)
            size_kb = round(os.path.getsize(file_path) / 1024, 2)
            result += f"  • {file} ({size_kb} KB)\n"
        
        return result.strip()
    except Exception as e:
        return f"Error listing documents: {str(e)}"

@mcp.tool()
def copy_document(source_filename: str, target_filename: str) -> str:
    """Copy a Word document to create a new version while preserving all formatting.
    
    Args:
        source_filename: Path to the source Word document to copy
        target_filename: Path for the new copied document
    
    Returns:
        Success or error message
    """
    source_filename = ensure_docx_extension(source_filename)
    target_filename = ensure_docx_extension(target_filename)
    
    if not check_file_exists(source_filename):
        return f"Error: Source document '{source_filename}' does not exist"
    
    if check_file_exists(target_filename):
        return f"Error: Target document '{target_filename}' already exists"
    
    try:
        # Load the source document
        source_doc = Document(source_filename)
        
        # Save it with the new filename (this preserves all formatting, styles, etc.)
        source_doc.save(target_filename)
        
        return f"Document copied successfully from '{source_filename}' to '{target_filename}'"
    except Exception as e:
        return f"Error copying document: {str(e)}"

# WRITE OPERATIONS

@mcp.tool()
def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document.
    
    Args:
        filename: Name for the new document
        title: Optional title for document metadata
        author: Optional author for document metadata
    
    Returns:
        Success or error message
    """
    filename = ensure_docx_extension(filename)
    
    try:
        doc = Document()
        
        # Set metadata if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        doc.save(filename)
        return f"Document '{filename}' created successfully"
    except Exception as e:
        return f"Error creating document: {str(e)}"

@mcp.tool()
def write_text(filename: str, text: str, append: bool = True) -> str:
    """Write text to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Text content to write
        append: If True, append to existing content. If False, replace all content.
    
    Returns:
        Success or error message
    """
    filename = ensure_docx_extension(filename)
    
    try:
        if append and check_file_exists(filename):
            # Append to existing document
            doc = Document(filename)
        else:
            # Create new document or replace existing
            doc = Document()
        
        # Add text as paragraph
        doc.add_paragraph(text)
        doc.save(filename)
        
        action = "appended to" if append and check_file_exists(filename) else "written to"
        return f"Text {action} '{filename}' successfully"
    except Exception as e:
        return f"Error writing to document: {str(e)}"

@mcp.tool()
def add_heading(filename: str, text: str, level: int = 1) -> str:
    """Add a heading to a Word document.
    
    Args:
        filename: Path to the Word document
        text: Heading text
        level: Heading level (1-6, where 1 is largest)
    
    Returns:
        Success or error message
    """
    filename = ensure_docx_extension(filename)
    
    # Validate heading level
    if level < 1 or level > 6:
        return "Error: Heading level must be between 1 and 6"
    
    try:
        if check_file_exists(filename):
            doc = Document(filename)
        else:
            doc = Document()
        
        doc.add_heading(text, level=level)
        doc.save(filename)
        
        return f"Heading '{text}' (level {level}) added to '{filename}'"
    except Exception as e:
        return f"Error adding heading: {str(e)}"

@mcp.tool()
def replace_text(filename: str, find_text: str, replace_text: str) -> str:
    """Find and replace text in a Word document.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
    
    Returns:
        Number of replacements made or error message
    """
    filename = ensure_docx_extension(filename)
    
    if not check_file_exists(filename):
        return f"Error: Document '{filename}' does not exist"
    
    try:
        doc = Document(filename)
        replacements = 0
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                # Replace in runs to preserve formatting
                for run in paragraph.runs:
                    if find_text in run.text:
                        run.text = run.text.replace(find_text, replace_text)
                        replacements += run.text.count(replace_text)
        
        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            for run in paragraph.runs:
                                if find_text in run.text:
                                    run.text = run.text.replace(find_text, replace_text)
                                    replacements += run.text.count(replace_text)
        
        if replacements > 0:
            doc.save(filename)
            return f"Replaced {replacements} occurrence(s) of '{find_text}' with '{replace_text}'"
        else:
            return f"No occurrences of '{find_text}' found"
    except Exception as e:
        return f"Error replacing text: {str(e)}"

@ mcp.tool()
def export_to_pdf(source_filename: str, target_filename: Optional[str] = None) -> str:
    """Export a Word document to PDF format.
    
    Args:
        source_filename: Path to the source Word document
        target_filename: Optional path for the PDF output (defaults to same name with .pdf extension)
    
    Returns:
        Success or error message
    
    Note:
        Requires LibreOffice to be installed and available in PATH as 'soffice'
    """
    source_path = ensure_docx_extension(source_filename)
    if not check_file_exists(source_path):
        return f"Error: Source document '{source_path}' does not exist"
    
    # Determine target path
    if target_filename is None:
        target_path = source_path.replace('.docx', '.pdf')
    else:
        target_path = target_filename
        if not target_path.lower().endswith('.pdf'):
            target_path += '.pdf'
    
    output_dir = os.path.dirname(target_path) or '.'
    
    try:
        # Run LibreOffice conversion
        cmd = [
            'soffice',
            '--headless',
            '--convert-to', 'pdf',
            source_path,
            '--outdir', output_dir
        ]
        result = subprocess.run(cmd, capture_output=True, text=True)
        
        if result.returncode != 0:
            error_msg = result.stderr.strip() or result.stdout.strip()
            return f"Conversion failed: {error_msg}"
        
        # Determine the converted filename (same base as source)
        source_base = os.path.splitext(os.path.basename(source_path))[0]
        converted_path = os.path.join(output_dir, f"{source_base}.pdf")
        
        # If target_path is different, rename the file
        if converted_path != target_path:
            if os.path.exists(target_path):
                return f"Error: Target file '{target_path}' already exists"
            os.rename(converted_path, target_path)
        
        return f"Document converted to PDF: '{target_path}'"
    
    except Exception as e:
        return f"Error during PDF conversion: {str(e)}"

def main():
    """Main entry point for the server."""
    print("Starting Simple Word Document MCP Server...")
    print("Available tools:")
    print("  READ: read_document, get_document_info, list_documents")
    print("  COPY: copy_document")
    print("  WRITE: create_document, write_text, add_heading, replace_text, export_to_pdf")
    print()
    
    try:
        # Run with stdio transport (default for MCP)
        mcp.run(transport='stdio')
    except KeyboardInterrupt:
        print("\nShutting down server...")
    except Exception as e:
        print(f"Error starting server: {e}")

if __name__ == "__main__":
    main()